import customtkinter as ctk
import tkinter as tk
from tkinter import simpledialog
from tkinter import filedialog
import docx  # Import the docx module

# Initialize an empty list to store item data
items = []

# Function to add an item
def add_item():
    item_name = entry_name.get()
    item_percentage = entry_percentage.get()

    
    if item_name and item_percentage:
        items.append((item_name, item_percentage))
        entry_name.delete(0, ctk.END)
        entry_percentage.delete(0, ctk.END)
        entry_name.focus()
    else:
        tk.messagebox.showwarning("Warning", "Both name and percentage must be filled.")
    
    print(items)


# Function to generate a document
def generate_document():
    if not items:
        ctk.messagebox.showwarning("Warning", "No items to generate a document.")
        return

    file_path = filedialog.asksaveasfilename(defaultextension=".docx", filetypes=[("Word Documents", "*.docx")])

    if file_path:
            # Create a new Word document
            doc = docx.Document()

            # Add a title to the document
            doc.add_heading("RESULTS", level=1)

            # Add items from the 'items' list to the document
            for item_name, item_percentage in items:
                doc.add_paragraph(f"{item_name}{'.'*(20-len(item_name))}{item_percentage}%")

            # Save the document to the specified file path
            doc.save(file_path)
            ctk.messagebox.showinfo("Success", "Document generated and saved successfully.")
            root.destroy()

# Create the main window
root = ctk.CTk()
root.title("Item List Generator")
root.resizable(False,False)

# Create and place widgets
label_name = ctk.CTkLabel(root, text="Input name:")
label_percentage = ctk.CTkLabel(root, text="Input percentage:")

entry_name = ctk.CTkEntry(root)
entry_name.focus()
entry_percentage = ctk.CTkEntry(root)

button_add_item = ctk.CTkButton(root, text="Add Another Item", command=add_item)
button_generate_document = ctk.CTkButton(root, text="Generate Document", command=generate_document)

# Layout widgets
label_name.grid(row=0, column=0, padx=4, pady=5)
label_percentage.grid(row=1, column=0, padx=4, pady=5)
entry_name.grid(row=0, column=1, padx=4, pady=5)
entry_percentage.grid(row=1, column=1, padx=4, pady=5)
button_add_item.grid(row=2, column=0, padx=4, pady=5)
button_generate_document.grid(row=2, column=1, padx=4, pady=5)

root.mainloop()

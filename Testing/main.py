def add_two_number(num1,num2):
    return num1 + num2

def cal_length(chemic, symbol):
    length = (74 - len(chemic) + len(symbol)) * '.'
    
    print(length)

periodic_table = {
    "Hydrogen": "H",
    "Helium": "He",
    "Lithium": "Li",
    "Beryllium": "Be",
    "Boron": "B",
    "Carbon": "C",
    "Nitrogen": "N",
    "Oxygen": "O",
    "Fluorine": "F",
    "Neon": "Ne",
    # Add more elements and symbols as needed
}
# for element,sym in periodic_table.items():
    # cal_length(element,sym)
from docx import Document
from docx.shared import Inches

document = Document()

p = document.add_paragraph()
r = p.add_run()
r.add_text('Good Morning every body,This is my ')
r.add_picture('/home/saabir/Music/Python-Projects/Testing/rock.jpg')
r.add_text(' do you like it?')

document.save('demo.docx')
